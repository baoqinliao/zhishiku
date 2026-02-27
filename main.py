import base64
import gc
import os
import shutil
import sys
from io import BytesIO

import lxml.html
import pandas as pd
from langchain_core.documents import Document
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from llama_cpp import Llama
from llama_cpp.llama_chat_format import Llava15ChatHandler

from pydantic import BaseModel
# --- 核心：替换为 RapidOCR ---
from rapidocr_onnxruntime import RapidOCR
# 核心 AI 库
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.llms import LlamaCpp
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain, RetrievalQA
from langchain_community.tools import DuckDuckGoSearchRun
from langchain.agents import initialize_agent, Tool, AgentType
from langchain_core.prompts import ChatPromptTemplate




from langchain.chains import create_history_aware_retriever
from langchain_core.prompts import MessagesPlaceholder
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory

# --- 2. 记忆管理逻辑 ---
store = {}


def get_session_history(session_id: str):
    """获取并裁剪历史记录，防止上下文溢出"""
    if session_id not in store:
        store[session_id] = ChatMessageHistory()

    # 【核心优化】：只保留最近 10 条消息 (约 5 轮对话)
    if len(store[session_id].messages) > 10:
        store[session_id].messages = store[session_id].messages[-10:]

    return store[session_id]


# --- 新增配置 ---
LLAVA_MODEL_PATH = "./models/ggml-model-q4_k.gguf" # LLaVA 主模型
MM_PROJ_PATH = "./models/mmproj-model-f16.gguf"

# 全局初始化 OCR (首次启动会自动下模型，之后断网可用)
# 初始化 RapidOCR 引擎 (无 Paddle 依赖，兼容性极佳)
ocr_engine = RapidOCR()



_global_llava_instance = None

def get_llava_instance():
    global _global_llava_instance
    if _global_llava_instance is None:
        print("--- 正在加载 LLaVA 视觉模型 (仅加载一次) ---")
        chat_handler = Llava15ChatHandler(clip_model_path=MM_PROJ_PATH)
        _global_llava_instance = Llama(
            model_path=LLAVA_MODEL_PATH,
            chat_handler=chat_handler,
            n_ctx=4096,  # 确保足够长
            n_threads=8,
            n_gpu_layers=-1, # 如果有显卡建议开启，否则设为 0
            verbose=False
        )
    return _global_llava_instance




def process_image_to_text(file_path, custom_prompt=None):
    # 1. OCR 提取（这步通常没问题）
    result, _ = ocr_engine(file_path)
    ocr_text = "\n".join([line[1] for line in result]) if result else ""

    # 2. 准备图片数据
    # with open(file_path, "rb") as f:
    #     # 修复点：确保 Base64 编码不包含换行符
    #     base64_img = base64.b64encode(f.read()).decode('utf-8').replace('\n', '')
    #     # 明确指定 mime type
    #     data_url = f"data:image/jpeg;base64,{base64_img}"

    # --- 修复点：Base64 处理逻辑 ---
    with open(file_path, "rb") as f:
        # llama-cpp-python 有时更喜欢纯 base64，不带 data: 前缀
        # 或者使用它自带的工具函数。这里我们尝试最稳妥的 Data URL 格式
        image_bytes = f.read()
        base64_img = base64.b64encode(image_bytes).decode('utf-8')
        data_url = f"data:image/jpeg;base64,{base64_img}"



    llava = get_llava_instance()

    prompt = custom_prompt if custom_prompt else "详细描述这张图里的内容、颜色、材质和文字。"

    try:
        # 增加超时和异常捕获
        response = llava.create_chat_completion(
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}}  # 注意这里的嵌套结构
                ]
            }],
            max_tokens=1024 # 适当增加输出长度
        )
        visual_desc = response["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"LLaVA 分析核心报错: {e}")
        visual_desc = "图片视觉解析失败，请检查模型配置。"


    return f"【语义】：{visual_desc}\n【文字】：{ocr_text}"






app = FastAPI(title="智识库 Pro")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# --- 配置 (参考你的原始参数) ---
MODEL_PATH = "./models/Qwen2.5-7b-instruct-q4_k_m.gguf"
EMBEDDING_MODEL_PATH = r"D:\models\paraphrase-multilingual-MiniLM-L12-v2"
DB_DIR = "./chroma_db"

# --- 初始化持久化资源 ---
print("--- 智识库系统启动中 ---")
embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_PATH)
# 建立持久化向量数据库
if not os.path.exists(DB_DIR): os.makedirs(DB_DIR)
vector_db = Chroma(persist_directory=DB_DIR, embedding_function=embeddings)

# 加载 LlamaCpp
llm = LlamaCpp(
    model_path=MODEL_PATH,
    n_ctx=8192,
    n_threads=8,  # 匹配 4800H 核心
    streaming=True,
    verbose=False,
    temperature=0.1,
    stop=["<|im_start|>", "<|im_end|>", "用户:", "问：", "Human:", "###"],
    repeat_penalty=1.2
)

search_tool = DuckDuckGoSearchRun()

class ChatRequest(BaseModel):
    query: str
    use_online: bool = False
    session_id: str = "default_user"  # 新增：用于区分不同会话

# --- 提示词定义 ---
# system_prompt_str = (
#     "你是一个基于本地资料的问答助手。你的任务是根据【背景资料】简洁回答问题。\n\n"
#     "【背景资料】：\n{context}\n\n"
#     "【要求】：直接给出答案，不要重复背景资料的标题。本地资料如果没有，请不要发散，直接回答本地资料库没有"
# )
# prompt_template = ChatPromptTemplate.from_messages([
#     ("system", system_prompt_str),
#     ("human", "{input}"),
# ])


# class ChatRequest(BaseModel):
#     query: str
#     use_online: bool = False


# --- API 接口 ---

# 尝试导入 python-docx
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

class DocxLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        if not DocxDocument:
            print("未安装 python-docx，无法解析 Word 文档")
            return []
        try:
            doc = DocxDocument(self.file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))

            return [Document(page_content="\n".join(full_text), metadata={"source": self.file_path})]
        except Exception as e:
            print(f"Word解析失败: {e}")
            return []

class ExcelLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        try:
            # 读取所有工作表
            with pd.ExcelFile(self.file_path) as xls:
                documents = []
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    # 将每一行转换为文本格式: "列名: 值, 列名: 值"
                    text_rows = []
                    for _, row in df.iterrows():
                        row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                        text_rows.append(row_str)

                    content = f"Sheet: {sheet_name}\n" + "\n".join(text_rows)
                    documents.append(Document(page_content=content, metadata={"source": self.file_path, "sheet": sheet_name}))
                return documents
        except Exception as e:
            print(f"Excel解析失败: {e}")
            return []

class HTMLLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        try:
            with open(self.file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            tree = lxml.html.fromstring(html_content)
            text = tree.text_content()
            # 简单的空白清洗
            text = " ".join(text.split())
            return [Document(page_content=text, metadata={"source": self.file_path})]
        except Exception as e:
            print(f"HTML解析失败: {e}")
            return []

@app.get("/", response_class=HTMLResponse)
async def get_index():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    """上传并自动学习文档"""
    filename = file.filename
    temp_path = f"./static/{filename}"
    if not os.path.exists("./static"): os.makedirs("./static")
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:

        docs = []
        lower_name = filename.lower()
        # 支持 PDF, TXT, MD, HTML
        # 1. 处理图片 (jpg, png)
        if lower_name.endswith((".jpg", ".jpeg", ".png")):
            content = process_image_to_text(temp_path)
            # 存入向量库时，一定要在 metadata 里记下图片路径，方便前端回显图片
            docs = [Document(page_content=content, metadata={"source": temp_path, "type": "image"})]

        # 2. 处理 PDF (增加 OCR 兜底)
        elif lower_name.endswith(".pdf"):
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
            # 扫描件判定
            if not docs or len("".join([d.page_content for d in docs]).strip()) < 10:
                from pdf2image import convert_from_path
                images = convert_from_path(temp_path)
                ocr_text_all = ""
                for i, image in enumerate(images):
                    img_buf = BytesIO()
                    image.save(img_buf, format='JPEG')
                    # 调用 RapidOCR
                    res, _ = ocr_engine(img_buf.getvalue())
                    if res:
                        ocr_text_all += f"\n--- 第 {i + 1} 页 ---\n" + "\n".join([l[1] for l in res])
                docs = [Document(page_content=ocr_text_all, metadata={"source": temp_path, "type": "scanned_pdf"})]


        # 3. 其他原有逻辑 (Excel, Word...)
        elif lower_name.endswith(".docx"):
            docs = DocxLoader(temp_path).load()
        elif lower_name.endswith((".xlsx", ".xls")):
            docs = ExcelLoader(temp_path).load()
        else:
            docs = TextLoader(temp_path, encoding='utf-8').load()

        if not docs:
            return {"status": "error", "message": "无法解析内容"}

        # 切分并入库
        # splits = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=100).split_documents(docs)
        # vector_db.add_documents(splits)
        splits = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=80).split_documents(docs)
        vector_db.add_documents(splits)
        return {"status": "success", "filename": filename}

    finally:
        # 注意：图片建议保留在 static 文件夹，不要删除！
        # 这样对话时返回 metadata 里的路径，前端才能 <img> 展示给客户看

        # 修复点 2：finally 块只负责删除非图片文件，不再包含解析逻辑
        if not lower_name.endswith((".jpg", ".png", ".jpeg")):
            if os.path.exists(temp_path): os.remove(temp_path)

        # if not lower_name.endswith((".jpg", ".png", ".jpeg")):
        #     if os.path.exists(temp_path): os.remove(temp_path)
        # docs = loader.load()
        # if not docs:
        #      return {"status": "error", "message": "无法读取文件内容"}
        #
        # splits = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=40).split_documents(docs)
        # # 增量添加到本地数据库
        # vector_db.add_documents(splits)
        # return {"status": "success"}



@app.get("/files")
async def get_files():
    """从数据库元数据中提取唯一文件名列表"""
    data = vector_db.get()
    if not data or not data["metadatas"]:
        return {"files": []}
    # 提取 source 字段并去重
    sources = list(set([os.path.basename(m.get("source", "未知文档")) for m in data["metadatas"]]))
    return {"files": sources}


@app.post("/chat")
async def chat(request: ChatRequest):
    # --- A. 无论哪种模式，都先进行“问题重写”以支持多轮对话 ---
    contextualize_q_system_prompt = (
        "根据聊天记录和最新的用户提问，如果提问涉及之前的上下文，"
        "请将其重写为一个独立的、明确的问题。如果不需要重写，请原样输出。不要回答问题。"
    )
    contextualize_q_prompt = ChatPromptTemplate.from_messages([
        ("system", contextualize_q_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ])

    # 这是一个不带检索的小链条，专门用来把“它有什么用”变成“Python 有什么用”
    rewrite_chain = contextualize_q_prompt | llm

    # 获取历史并生成重写后的问题
    history = get_session_history(request.session_id)
    rewritten_query = rewrite_chain.invoke({"input": request.query, "chat_history": history.messages})
    # 清洗掉可能生成的标签
    if "Assistant:" in rewritten_query: rewritten_query = rewritten_query.split("Assistant:")[-1].strip()


    if not request.use_online:

        # 2. 检索器封装
        retriever = vector_db.as_retriever(
            search_type="mmr",
            search_kwargs={
                "k": 10,  # 针对精准信息，k 不需要太大，5 足够
                "fetch_k": 30,  # 扩大初步筛选范围
                "lambda_mult": 0.8  # 调高相关性权重，降低多样性干扰
            }
        )
        history_aware_retriever = create_history_aware_retriever(llm, retriever, contextualize_q_prompt)

        qa_system_prompt = (
                "你是一个基于本地资料的问答助手。你的任务是根据【背景资料】简洁回答问题。\n\n"
    "【背景资料】：\n{context}\n\n"
    "【要求】：直接给出答案，不要重复背景资料的标题。"
        )



        qa_prompt = ChatPromptTemplate.from_messages([
            ("system", qa_system_prompt),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
        ])

        qa_chain = create_stuff_documents_chain(llm, qa_prompt)
        rag_chain = create_retrieval_chain(history_aware_retriever, qa_chain)

        conversational_rag_chain = RunnableWithMessageHistory(
            rag_chain, get_session_history,
            input_messages_key="input", history_messages_key="chat_history", output_messages_key="answer"
        )

        res = conversational_rag_chain.invoke(
            {"input": request.query},
            config={"configurable": {"session_id": request.session_id}}
        )

        # 提取图片/来源
        source_images = []
        for doc in res.get("context", []):
            if doc.metadata.get("type") == "image" or doc.metadata.get("source", "").lower().endswith(
                    (".jpg", ".png", ".jpeg")):
                img_url = f"/static/{os.path.basename(doc.metadata.get('source'))}"
                if img_url not in source_images: source_images.append(img_url)

        final_answer = res["answer"]
        if "Assistant:" in final_answer: final_answer = final_answer.split("Assistant:")[-1].strip()

        return {
            "answer": final_answer,
            "mode": "本地库模式",
            "images": source_images,
            "source_docs": list(set([os.path.basename(d.metadata.get("source")) for d in res.get("context", [])]))
        }

    else:

        search_result = search_tool.run(rewritten_query)
        qa_system_prompt = (
            "你是一个联网助手。请根据以下【联网搜索结果】回答用户的问题。\n\n"
            "【搜索结果】：\n{context}"
        )
        qa_prompt = ChatPromptTemplate.from_messages([
            ("system", qa_system_prompt),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
        ])

        # 构建问答链
        chain = qa_prompt | llm
        full_res = chain.invoke({"input": request.query, "chat_history": history.messages, "context": search_result})

        # 更新记忆
        history.add_user_message(request.query)
        history.add_ai_message(full_res)

        return {
            "answer": full_res.split("Assistant:")[-1].strip() if "Assistant:" in full_res else full_res.strip(),
            "mode": "联网增强模式",
            "rewritten_query": rewritten_query
        }





@app.delete("/files/{filename}")
async def delete_file(filename: str):
    """删除指定文件"""
    # 构建可能的路径 (取决于入库时的路径格式)
    target_source = f"./static/{filename}"

    # 直接操作 Chroma 底层 Collection 进行删除
    try:
        vector_db._collection.delete(where={"source": target_source})
        return {"status": "deleted", "file": filename}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.delete("/clear")
async def clear():
    global vector_db
    # 1. 彻底销毁对象并清空目录
    vector_db = None
    gc.collect()
    if os.path.exists(DB_DIR): shutil.rmtree(DB_DIR)
    os.makedirs(DB_DIR)

    # 2. 清理 static 目录下的非 UI 文件（只删图片和文档）
    ui_files = ['index.html', 'vue.global.js', 'tailwind.min.js', 'all.min.css']
    for f in os.listdir("./static"):
        if f not in ui_files and os.path.isfile(os.path.join("./static", f)):
            try:
                os.remove(os.path.join("./static", f))
            except:
                pass

    # 3. 重新初始化
    vector_db = Chroma(persist_directory=DB_DIR, embedding_function=embeddings)
    store.clear()
    return {"status": "cleared"}


# 挂载本地静态资源文件夹
app.mount("/static", StaticFiles(directory="static"), name="static")

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)