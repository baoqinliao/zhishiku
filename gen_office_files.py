import pandas as pd
from docx import Document

def create_sample_docx():
    doc = Document()
    doc.add_heading('项目合同书 (测试样本)', 0)

    doc.add_paragraph('甲方：未来科技有限公司')
    doc.add_paragraph('乙方：智识库软件服务商')

    doc.add_heading('第一条 项目内容', level=1)
    doc.add_paragraph(
        '乙方负责为甲方开发“企业级本地知识库系统”，包括但不限于以下功能：'
        '文档解析、向量存储、自然语言问答。'
    )

    doc.add_heading('第二条 费用及支付', level=1)
    
    # 添加表格
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '阶段'
    hdr_cells[1].text = '金额 (万元)'
    hdr_cells[2].text = '付款条件'

    data = [
        ('首付款', '10.0', '合同签订后5个工作日内'),
        ('中期款', '15.0', '系统上线试运行'),
        ('尾款', '5.0', '验收合格后'),
    ]

    for stage, amount, condition in data:
        row_cells = table.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = amount
        row_cells[2].text = condition

    doc.save('sample_contract.docx')
    print("生成完成: sample_contract.docx")

def create_sample_xlsx():
    data = {
        '产品名称': ['高性能服务器', '企业级交换机', '防火墙', 'UPS电源'],
        '型号': ['SRV-2025', 'SW-48G', 'FW-X1', 'UPS-3000'],
        '库存数量': [15, 40, 8, 20],
        '单价 (元)': [25000, 8000, 12000, 3500],
        '仓库位置': ['A区-01', 'B区-05', 'A区-02', 'C区-10']
    }
    
    df = pd.DataFrame(data)
    df.to_excel('sample_inventory.xlsx', index=False)
    print("生成完成: sample_inventory.xlsx")

if __name__ == "__main__":
    try:
        create_sample_docx()
        create_sample_xlsx()
        print("所有测试文件已生成。")
    except Exception as e:
        print(f"生成失败: {e}")
        print("请确保已安装 pandas, openpyxl, python-docx")
