import os
import sys
from pathlib import Path

import pandas as pd
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Optional dependencies check
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import pypdf
except ImportError:
    pypdf = None

def clean_text(text):
    """
    清洗文本：去除多余空白行和空格
    """
    if not text:
        return ""
    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    text = '\n'.join(chunk for chunk in chunks if chunk)
    return text

def process_image(file_path):
    """
    处理图片文件，使用 Tesseract OCR 提取文本
    """
    if not (Image and pytesseract):
        print(f"[WARN] 跳过图片 {file_path.name}: 请安装 'Pillow' 和 'pytesseract' (并确保系统安装了 Tesseract-OCR)。")
        return ""
    try:
        # 默认尝试识别中文简体和英文
        # 如果未安装中文语言包，可能需要改为 lang='eng'
        text = pytesseract.image_to_string(Image.open(file_path), lang='chi_sim+eng')
        return clean_text(text)
    except Exception as e:
        print(f"[ERROR] 图片处理失败 {file_path.name}: {e}")
        return ""

def process_html(file_path):
    """
    处理 HTML 文件，提取纯文本
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        if BeautifulSoup:
            soup = BeautifulSoup(content, 'html.parser')
            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.extract()
            text = soup.get_text()
        else:
            # 简单回退方案
            import re
            text = re.sub('<[^<]+?>', '', content)
            
        return clean_text(text)
    except Exception as e:
        print(f"[ERROR] HTML 处理失败 {file_path.name}: {e}")
        return ""

def process_pdf(file_path):
    """
    处理 PDF 文件
    """
    if not pypdf:
        print(f"[WARN] 跳过 PDF {file_path.name}: 未安装 'pypdf'。")
        return ""
    try:
        reader = pypdf.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return clean_text(text)
    except Exception as e:
        print(f"[ERROR] PDF 处理失败 {file_path.name}: {e}")
        return ""

def process_docx(file_path):
    """
    处理 Word 文档
    """
    if not DocxDocument:
        print(f"[WARN] 跳过 Word {file_path.name}: 未安装 'python-docx'。")
        return ""
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return clean_text("\n".join(full_text))
    except Exception as e:
        print(f"[ERROR] Word 处理失败 {file_path.name}: {e}")
        return ""

def process_excel(file_path):
    """
    处理 Excel 表格
    """
    try:
        xls = pd.ExcelFile(file_path)
        full_text = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            full_text.append(f"Sheet: {sheet_name}")
            for _, row in df.iterrows():
                row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                full_text.append(row_str)
        return clean_text("\n".join(full_text))
    except Exception as e:
        print(f"[ERROR] Excel 处理失败 {file_path.name}: {e}")
        return ""

def process_text(file_path):
    """
    处理普通文本和 Markdown
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return clean_text(f.read())
    except Exception as e:
        print(f"[ERROR] 文本读取失败 {file_path.name}: {e}")
        return ""

def main():
    if len(sys.argv) < 3:
        print("用法: python data_processor.py <输入目录> <输出目录>")
        print("示例: python data_processor.py ./raw_data ./clean_data")
        print("\n说明: 本脚本将递归扫描输入目录，将 PDF, HTML, MD, 图片等转换为纯文本并保存到输出目录。")
        return

    input_dir = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])

    if not input_dir.exists():
        print(f"错误: 输入目录 '{input_dir}' 不存在。")
        return

    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"--- 开始处理数据: '{input_dir}' -> '{output_dir}' ---")
    
    count = 0
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            file_path = Path(root) / file
            ext = file_path.suffix.lower()
            
            content = ""
            if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                content = process_image(file_path)
            elif ext in ['.html', '.htm']:
                content = process_html(file_path)
            elif ext == '.pdf':
                content = process_pdf(file_path)
            elif ext in ['.txt', '.md', '.markdown']:
                content = process_text(file_path)
            elif ext == '.docx':
                content = process_docx(file_path)
            elif ext in ['.xlsx', '.xls']:
                content = process_excel(file_path)
            else:
                # 跳过不支持的文件类型
                continue

            if content:
                # 生成输出文件名，保留原文件名但改为 .txt 后缀
                out_name = f"{file_path.stem}_{ext[1:]}_cleaned.txt"
                out_path = output_dir / out_name
                
                with open(out_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"[成功] 处理: {file} -> {out_name}")
                count += 1
            else:
                print(f"[跳过] 空内容或处理失败: {file}")

    print(f"--- 处理完成! 共生成 {count} 个清洗后的文档 ---")

if __name__ == "__main__":
    main()
