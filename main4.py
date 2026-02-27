import os
import shutil
import sys
import lxml.html
import pandas as pd
from langchain_core.documents import Document
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# 核心 AI 库
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.llms import LlamaCpp
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain, RetrievalQA
from langchain_community.tools import DuckDuckGoSearchRun
from langchain.agents import initialize_agent, Tool, AgentType
from langchain_core.prompts import ChatPromptTemplate

app = FastAPI(title="智识库 Pro")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# --- 配置 (参考你的原始参数) ---
MODEL_PATH = "./models/Qwen2.5-7b-instruct-q4_k_m.gguf"
EMBEDDING_MODEL_PATH = r"D:\models\paraphrase-multilingual-MiniLM-L12-v2"
DB_DIR = "./chroma_db"

# --- 初始化持久化资源 ---
print("--- 智识库系统启动中 ---")
embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_PATH)
# 建立持久化向量数据库
if not os.path.exists(DB_DIR): os.makedirs(DB_DIR)
vector_db = Chroma(persist_directory=DB_DIR, embedding_function=embeddings)

# 加载 LlamaCpp
llm = LlamaCpp(
    model_path=MODEL_PATH,
    n_ctx=8192,
    n_threads=8,  # 匹配 4800H 核心
    streaming=True,
    verbose=False,
    temperature=0.1,
    stop=["<|im_start|>", "<|im_end|>", "用户:", "问：", "Human:", "###"],
    repeat_penalty=1.2
)

search_tool = DuckDuckGoSearchRun()

# --- 提示词定义 ---
system_prompt_str = (
    "你是一个基于本地资料的问答助手。你的任务是根据【背景资料】简洁回答问题。\n\n"
    "【背景资料】：\n{context}\n\n"
    "【要求】：直接给出答案，不要重复背景资料的标题。"
)
prompt_template = ChatPromptTemplate.from_messages([
    ("system", system_prompt_str),
    ("human", "{input}"),
])


class ChatRequest(BaseModel):
    query: str
    use_online: bool = False


# --- API 接口 ---

# 尝试导入 python-docx
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

class DocxLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        if not DocxDocument:
            print("未安装 python-docx，无法解析 Word 文档")
            return []
        try:
            doc = DocxDocument(self.file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            
            return [Document(page_content="\n".join(full_text), metadata={"source": self.file_path})]
        except Exception as e:
            print(f"Word解析失败: {e}")
            return []

class ExcelLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        try:
            # 读取所有工作表
            with pd.ExcelFile(self.file_path) as xls:
                documents = []
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    # 将每一行转换为文本格式: "列名: 值, 列名: 值"
                    text_rows = []
                    for _, row in df.iterrows():
                        row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                        text_rows.append(row_str)
                    
                    content = f"Sheet: {sheet_name}\n" + "\n".join(text_rows)
                    documents.append(Document(page_content=content, metadata={"source": self.file_path, "sheet": sheet_name}))
                return documents
        except Exception as e:
            print(f"Excel解析失败: {e}")
            return []

class HTMLLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        try:
            with open(self.file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            tree = lxml.html.fromstring(html_content)
            text = tree.text_content()
            # 简单的空白清洗
            text = " ".join(text.split())
            return [Document(page_content=text, metadata={"source": self.file_path})]
        except Exception as e:
            print(f"HTML解析失败: {e}")
            return []

@app.get("/", response_class=HTMLResponse)
async def get_index():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    """上传并自动学习文档"""
    temp_path = f"./static/{file.filename}"
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 支持 PDF, TXT, MD, HTML
        if temp_path.lower().endswith(".pdf"):
            loader = PyPDFLoader(temp_path)
        elif temp_path.lower().endswith(".html") or temp_path.lower().endswith(".htm"):
            loader = HTMLLoader(temp_path)
        elif temp_path.lower().endswith(".docx"):
            loader = DocxLoader(temp_path)
        elif temp_path.lower().endswith(".xlsx") or temp_path.lower().endswith(".xls"):
            loader = ExcelLoader(temp_path)
        else:
            # 默认为文本处理 (txt, md)
            loader = TextLoader(temp_path, encoding='utf-8')

        docs = loader.load()
        if not docs:
             return {"status": "error", "message": "无法读取文件内容"}

        splits = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=40).split_documents(docs)
        # 增量添加到本地数据库
        vector_db.add_documents(splits)
        return {"status": "success"}
    finally:
        if os.path.exists(temp_path): os.remove(temp_path)


@app.get("/files")
async def get_files():
    """从数据库元数据中提取唯一文件名列表"""
    data = vector_db.get()
    if not data or not data["metadatas"]:
        return {"files": []}
    # 提取 source 字段并去重
    sources = list(set([os.path.basename(m.get("source", "未知文档")) for m in data["metadatas"]]))
    return {"files": sources}


@app.post("/chat")
async def chat(request: ChatRequest):
    """问答：隐私/联网双模式"""
    combine_docs_chain = create_stuff_documents_chain(llm, prompt_template)
    # qa_chain = create_retrieval_chain(vector_db.as_retriever(), combine_docs_chain)

    # 增加检索数量 k=10，确保关键信息不被挤出
    # retriever = vector_db.as_retriever(search_kwargs={"k": 10})
    retriever = vector_db.as_retriever(
        search_type="mmr",
        search_kwargs={
            "k": 10,  # 最终返回给大模型的文档数量
            "fetch_k": 20,  # 初步从数据库检索的文档数量（从中筛选出最相关的10个）
            "lambda_mult": 0.5  # 多样性权重：0.5 平衡相关性和多样性
        }
    )
    qa_chain = create_retrieval_chain(retriever, combine_docs_chain)

    if not request.use_online:
        res = qa_chain.invoke({"input": request.query})

        # --- 调试日志：打印检索到的文档片段 ---
        print(f"\n[DEBUG] 检索词: {request.query}")
        print(f"[DEBUG] 检索到的文档数量: {len(res.get('context', []))}")
        for i, doc in enumerate(res.get("context", [])):
            # 打印每个文档的前100个字符和来源
            source = doc.metadata.get("source", "未知来源")
            content_preview = doc.page_content[:100].replace('\n', ' ')
            print(f"[DEBUG] Doc {i} ({os.path.basename(source)}): {content_preview}...")
        print("-" * 50)
        # -----------------------------------

        full_answer = res["answer"]

        # 【新增清洗逻辑】：只保留 "Assistant:" 之后的内容
        if "Assistant:" in full_answer:
            final_answer = full_answer.split("Assistant:")[-1].strip()
        else:
            final_answer = full_answer.strip()

        return {"answer": final_answer, "mode": "本地库模式"}


        #
        # return {"answer": res["answer"], "mode": "本地库模式"}
    else:
        retrieval_tool = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=vector_db.as_retriever())
        tools = [
            Tool(name="本地知识库", func=retrieval_tool.run, description="查找私有信息"),
            Tool(name="联网搜索", func=search_tool.run, description="查找实时新闻")
        ]
        agent = initialize_agent(tools, llm, agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION, verbose=True)
        res = agent.run(request.query)
        return {"answer": res, "mode": "联网增强模式"}


@app.delete("/files/{filename}")
async def delete_file(filename: str):
    """删除指定文件"""
    # 构建可能的路径 (取决于入库时的路径格式)
    target_source = f"./static/{filename}"
    
    # 直接操作 Chroma 底层 Collection 进行删除
    try:
        vector_db._collection.delete(where={"source": target_source})
        return {"status": "deleted", "file": filename}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.delete("/clear")
async def clear():
    """一键清空本地知识库"""
    if os.path.exists(DB_DIR):
        shutil.rmtree(DB_DIR)
        os.makedirs(DB_DIR)
        # 重新初始化全局对象
        global vector_db
        vector_db = Chroma(persist_directory=DB_DIR, embedding_function=embeddings)
    return {"status": "cleared"}


# 挂载本地静态资源文件夹
app.mount("/static", StaticFiles(directory="static"), name="static")

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)